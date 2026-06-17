from pathlib import Path
import csv
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Dev\cc_fraud_detection")
PKG = ROOT / "publication_package"
FIG_DIR = PKG / "figures"
TABLE_DIR = PKG / "tables"
REF_DIR = PKG / "references"
SUPP_DIR = PKG / "supplementary_materials"
VER_DIR = PKG / "verification"
DOCX_PATH = PKG / "manuscript.docx"

for folder in [FIG_DIR, TABLE_DIR, REF_DIR, SUPP_DIR, VER_DIR]:
    folder.mkdir(parents=True, exist_ok=True)


REFERENCES = [
    ("bolton2002", 'R. J. Bolton and D. J. Hand, "Statistical fraud detection: A review," Statistical Science, vol. 17, no. 3, pp. 235-255, 2002, doi: 10.1214/ss/1042727940.'),
    ("ngai2011", 'E. W. T. Ngai, Y. Hu, Y. H. Wong, Y. Chen, and X. Sun, "The application of data mining techniques in financial fraud detection: A classification framework and an academic review of literature," Decision Support Systems, vol. 50, no. 3, pp. 559-569, 2011, doi: 10.1016/j.dss.2010.08.006.'),
    ("bhattacharyya2011", 'S. Bhattacharyya, S. Jha, K. Tharakunnel, and J. C. Westland, "Data mining for credit card fraud: A comparative study," Decision Support Systems, vol. 50, no. 3, pp. 602-613, 2011, doi: 10.1016/j.dss.2010.08.008.'),
    ("whitrow2009", 'C. Whitrow, D. J. Hand, P. Juszczak, D. Weston, and N. M. Adams, "Transaction aggregation as a strategy for credit card fraud detection," Data Mining and Knowledge Discovery, vol. 18, no. 1, pp. 30-55, 2009, doi: 10.1007/s10618-008-0116-z.'),
    ("bahnsen2016", 'A. C. Bahnsen, D. Aouada, A. Stojanovic, and B. Ottersten, "Feature engineering strategies for credit card fraud detection," Expert Systems with Applications, vol. 51, pp. 134-142, 2016, doi: 10.1016/j.eswa.2015.12.030.'),
    ("dalpozzolo2015", 'A. Dal Pozzolo, O. Caelen, R. A. Johnson, and G. Bontempi, "Calibrating probability with undersampling for unbalanced classification," in Proc. IEEE Symposium Series on Computational Intelligence, 2015.'),
    ("jurgovsky2018", 'J. Jurgovsky et al., "Sequence classification for credit-card fraud detection," Expert Systems with Applications, vol. 100, pp. 234-245, 2018, doi: 10.1016/j.eswa.2018.01.037.'),
    ("phua2010", 'C. Phua, V. Lee, K. Smith, and R. Gayler, "A comprehensive survey of data mining-based fraud detection research," arXiv:1009.6119, 2010.'),
    ("chawla2002", 'N. V. Chawla, K. W. Bowyer, L. O. Hall, and W. P. Kegelmeyer, "SMOTE: Synthetic minority over-sampling technique," Journal of Artificial Intelligence Research, vol. 16, pp. 321-357, 2002, doi: 10.1613/jair.953.'),
    ("he2009", 'H. He and E. A. Garcia, "Learning from imbalanced data," IEEE Transactions on Knowledge and Data Engineering, vol. 21, no. 9, pp. 1263-1284, 2009, doi: 10.1109/TKDE.2008.239.'),
    ("japkowicz2002", 'N. Japkowicz and S. Stephen, "The class imbalance problem: A systematic study," Intelligent Data Analysis, vol. 6, no. 5, pp. 429-449, 2002.'),
    ("galar2012", 'M. Galar, A. Fernandez, E. Barrenechea, H. Bustince, and F. Herrera, "A review on ensembles for the class imbalance problem: Bagging-, boosting-, and hybrid-based approaches," IEEE Transactions on Systems, Man, and Cybernetics, Part C, vol. 42, no. 4, pp. 463-484, 2012, doi: 10.1109/TSMCC.2011.2161285.'),
    ("branco2016", 'P. Branco, L. Torgo, and R. P. Ribeiro, "A survey of predictive modeling on imbalanced domains," ACM Computing Surveys, vol. 49, no. 2, pp. 1-50, 2016, doi: 10.1145/2907070.'),
    ("elkan2001", 'C. Elkan, "The foundations of cost-sensitive learning," in Proc. International Joint Conference on Artificial Intelligence, 2001, pp. 973-978.'),
    ("davis2006", 'J. Davis and M. Goadrich, "The relationship between Precision-Recall and ROC curves," in Proc. International Conference on Machine Learning, 2006, pp. 233-240, doi: 10.1145/1143844.1143874.'),
    ("saito2015", 'T. Saito and M. Rehmsmeier, "The precision-recall plot is more informative than the ROC plot when evaluating binary classifiers on imbalanced datasets," PLOS ONE, vol. 10, no. 3, 2015, doi: 10.1371/journal.pone.0118432.'),
    ("fawcett2006", 'T. Fawcett, "An introduction to ROC analysis," Pattern Recognition Letters, vol. 27, no. 8, pp. 861-874, 2006, doi: 10.1016/j.patrec.2005.10.010.'),
    ("breiman2001", 'L. Breiman, "Random forests," Machine Learning, vol. 45, no. 1, pp. 5-32, 2001, doi: 10.1023/A:1010933404324.'),
    ("quinlan1986", 'J. R. Quinlan, "Induction of decision trees," Machine Learning, vol. 1, no. 1, pp. 81-106, 1986, doi: 10.1007/BF00116251.'),
    ("chen2016", 'T. Chen and C. Guestrin, "XGBoost: A scalable tree boosting system," in Proc. ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 2016, pp. 785-794, doi: 10.1145/2939672.2939785.'),
    ("prokhorenkova2018", 'L. Prokhorenkova, G. Gusev, A. Vorobev, A. V. Dorogush, and A. Gulin, "CatBoost: Unbiased boosting with categorical features," in Advances in Neural Information Processing Systems, 2018.'),
    ("pedregosa2011", 'F. Pedregosa et al., "Scikit-learn: Machine learning in Python," Journal of Machine Learning Research, vol. 12, pp. 2825-2830, 2011.'),
    ("bergstra2012", 'J. Bergstra and Y. Bengio, "Random search for hyper-parameter optimization," Journal of Machine Learning Research, vol. 13, pp. 281-305, 2012.'),
    ("kohavi1995", 'R. Kohavi, "A study of cross-validation and bootstrap for accuracy estimation and model selection," in Proc. International Joint Conference on Artificial Intelligence, 1995, pp. 1137-1145.'),
    ("cawley2010", 'G. C. Cawley and N. L. C. Talbot, "On over-fitting in model selection and subsequent selection bias in performance evaluation," Journal of Machine Learning Research, vol. 11, pp. 2079-2107, 2010.'),
    ("varma2006", 'S. Varma and R. Simon, "Bias in error estimation when using cross-validation for model selection," BMC Bioinformatics, vol. 7, article 91, 2006, doi: 10.1186/1471-2105-7-91.'),
    ("kaufman2012", 'S. Kaufman, S. Rosset, C. Perlich, and O. Stitelman, "Leakage in data mining: Formulation, detection, and avoidance," ACM Transactions on Knowledge Discovery from Data, vol. 6, no. 4, pp. 1-21, 2012, doi: 10.1145/2382577.2382579.'),
    ("kapoor2023", 'S. Kapoor and A. Narayanan, "Leakage and the reproducibility crisis in machine-learning-based science," Patterns, vol. 4, no. 9, 2023, doi: 10.1016/j.patter.2023.100804.'),
    ("guyon2003", 'I. Guyon and A. Elisseeff, "An introduction to variable and feature selection," Journal of Machine Learning Research, vol. 3, pp. 1157-1182, 2003.'),
    ("kuhn2013", 'M. Kuhn and K. Johnson, Applied Predictive Modeling. New York, NY, USA: Springer, 2013.'),
]


def write_references():
    bib_lines = []
    for key, text in REFERENCES:
        title_match = re.search(r'"(.+?)"', text)
        title = title_match.group(1) if title_match else key
        bib_lines.append(f"@misc{{{key},\n  title = {{{title}}},\n  note = {{{text}}}\n}}\n")
    (REF_DIR / "references.bib").write_text("\n".join(bib_lines), encoding="utf-8")
    (REF_DIR / "references_ieee.txt").write_text(
        "\n".join(f"[{i}] {text}" for i, (_, text) in enumerate(REFERENCES, 1)),
        encoding="utf-8",
    )


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_1 = OxmlElement("w:fldChar")
    fld_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_2 = OxmlElement("w:fldChar")
    fld_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_1)
    run._r.append(instr)
    run._r.append(fld_2)


def add_para(doc, text="", style=None, bold=False, italic=False, align=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
    if align is not None:
        p.alignment = align
    return p


def add_caption(doc, text, kind):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6 if kind == "figure" else 3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    return p


def add_table_from_csv(doc, csv_path, caption, widths):
    add_caption(doc, caption, "table")
    with open(csv_path, newline="", encoding="utf-8") as handle:
        rows = list(csv.reader(handle))
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(value)
            run.font.size = Pt(8.2 if len(rows[0]) >= 6 else 8.7)
            if row_idx == 0:
                run.bold = True
                set_cell_shading(cell, "F2F4F7")
            if col_idx > 0 and len(value) < 16:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_repeat_table_header(table.rows[0])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)
    return table


def add_figure(doc, fig_path, caption, width=6.1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(fig_path), width=Inches(width))
    add_caption(doc, caption, "figure")


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_page_number(section)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.08
    for name, size, color in [
        ("Heading 1", 14, "2E74B5"),
        ("Heading 2", 12, "2E74B5"),
        ("Heading 3", 11, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def build_docx():
    write_references()
    doc = Document()
    configure_styles(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Progressive Feature Engineering and Imbalanced Machine Learning for Credit Card Fraud Detection")
    run.bold = True
    run.font.size = Pt(17)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Author Name(s), Affiliation(s), Email Address(es)")
    run.italic = True
    run.font.size = Pt(10)

    add_para(doc, "Abstract", bold=True)
    abstract = (
        "Credit card fraud detection is an imbalanced binary classification problem in which minority-class performance is more informative than overall accuracy. "
        "This paper presents a controlled tabular machine learning study using a researcher-generated synthetic credit card transaction dataset containing 499,985 unique records, including 62,500 fraudulent and 437,485 legitimate transactions. "
        "Five progressively enriched dataset versions are evaluated, ranging from baseline transaction attributes to engineered amount, time, merchant-risk, card-risk, and composite-risk features. "
        "Logistic Regression, Decision Tree, Random Forest, XGBoost, and CatBoost classifiers are trained under no-SMOTE and SMOTE settings. "
        "The experimental design uses a stratified 60/20/20 train-validation-test split, 5-fold cross-validation on the training split, PR-AUC as the hyperparameter tuning metric, validation-based model-family selection, and untouched test-set reporting. "
        "At the default threshold, SMOTE improves fraud-class F1 across all dataset versions; the strongest default-threshold result is obtained by Random Forest on v3 with SMOTE, reaching precision 0.998525, recall 0.758080, fraud-class F1 0.861846, and PR-AUC 0.855235. "
        "After validation-based threshold optimization, the practical difference between SMOTE and no-SMOTE is largely removed; one high-F1 tuned operating point is the v5 no-SMOTE Random Forest at threshold 0.80, with precision 0.999156, recall 0.758000, and fraud-class F1 0.862030. "
        "The results show that progressive tabular feature engineering improves fraud-class performance, while conclusions are limited by the synthetic dataset, absence of external validation, and leakage-risk concerns around proxy risk indicators."
    )
    add_para(doc, abstract)
    add_para(
        doc,
        "Keywords - credit card fraud detection; imbalanced classification; feature engineering; SMOTE; Random Forest; XGBoost; CatBoost; PR-AUC.",
        bold=True,
    )

    doc.add_heading("I. Introduction", level=1)
    doc.add_heading("A. Background and Motivation", level=2)
    add_para(doc, "Credit card fraud detection is a high-impact binary classification problem in which fraudulent transactions are less frequent than legitimate transactions but carry disproportionate operational risk. Classical fraud systems have relied on expert rules, threshold policies, and manual risk signals [1], [2]. Machine learning methods extend these approaches by learning statistical regularities from transaction-level data [3]-[8].")
    add_para(doc, "Because fraud detection is imbalanced, accuracy alone is not an adequate evaluation target. This manuscript therefore emphasizes fraud-class precision, recall, fraud-class F1, ROC-AUC, and PR-AUC, with average precision used for hyperparameter tuning [9]-[17].")
    doc.add_heading("B. Research Gap", level=2)
    add_para(doc, "Prior fraud detection studies often compare classifiers or resampling strategies, but fewer controlled repository-level studies isolate the effect of progressive feature enrichment while keeping the underlying transaction population fixed. This study evaluates five dataset versions built from the same cleaned master population so that performance differences can be attributed to feature availability rather than sample-size changes.")
    doc.add_heading("C. Aim and Contributions", level=2)
    add_para(doc, "The aim is to evaluate whether progressive tabular feature engineering improves fraud-class performance under a leakage-aware experimental protocol. The study contributes a reproducible five-version dataset preparation workflow, a no-SMOTE versus SMOTE comparison, validation-based model selection, untouched test-set reporting, and validation-based threshold optimization. The study does not evaluate language-processing features, institutional deployment, or external generalization.")
    doc.add_heading("D. Paper Organization", level=2)
    add_para(doc, "Section II reviews related work. Section III describes dataset construction and validation. Section IV presents the methodology. Section V reports the experimental results. Section VI discusses interpretation and reproducibility. Section VII states threats to validity, and Section VIII concludes the paper.")

    doc.add_heading("II. Related Work", level=1)
    doc.add_heading("A. Traditional Rule-Based Fraud Detection", level=2)
    add_para(doc, "Rule-based fraud systems encode expert knowledge through manually specified thresholds and transaction patterns. Such systems are interpretable but require maintenance when transaction behavior or attack strategies change [1], [2].")
    doc.add_heading("B. Machine Learning for Credit Card Fraud Detection", level=2)
    add_para(doc, "Machine learning methods have been widely applied to card fraud detection, including linear classifiers, decision trees, random forests, gradient boosting, and sequence models [3]-[8]. In this repository, the final active experiment includes Logistic Regression, Decision Tree, Random Forest, XGBoost, and CatBoost.")
    doc.add_heading("C. Imbalanced Classification and SMOTE", level=2)
    add_para(doc, "Fraud detection commonly presents class imbalance. SMOTE and related methods create synthetic minority-class examples to reduce majority-class dominance during learning [9]-[14]. This study applies SMOTE only inside the training pipeline after preprocessing and within cross-validation folds.")
    doc.add_heading("D. Feature Engineering for Tabular Fraud Detection", level=2)
    add_para(doc, "Feature engineering can transform transaction attributes into stronger predictors, including amount-derived variables, temporal indicators, merchant descriptors, and risk summaries [5], [29], [30]. The present study evaluates progressive feature groups from v1 to v5 while holding the master population constant.")
    doc.add_heading("E. Evaluation Metrics for Imbalanced Fraud Detection", level=2)
    add_para(doc, "PR-AUC and fraud-class F1 are emphasized because ROC-AUC and accuracy can overstate performance when the minority class is operationally important [15]-[17]. The repository uses average precision for tuning and validation-based fraud-class F1 for model-family selection.")

    doc.add_heading("III. Dataset Construction and Validation", level=1)
    doc.add_heading("A. Dataset Generation and Cleaning", level=2)
    add_para(doc, "The study uses a researcher-generated synthetic transaction dataset. The active final master file is data/research_master_dataset.csv. The raw merged provenance file contains 849,999 rows and 350,014 exact duplicate rows; it is not treated as a sixth experiment. After cleaning and transaction-level de-duplication, the final master contains 499,985 unique transactions.")
    add_table_from_csv(doc, TABLE_DIR / "table1_dataset_summary.csv", "TABLE I. Dataset Summary", [2700, 6660])
    doc.add_heading("B. Dataset Audit and Integrity Checks", level=2)
    add_para(doc, "The final datasets contain no missing values, no exact duplicate rows, and no duplicate TransactionID values. Transaction identifiers are deterministic synthetic hashes used for reproducibility and split-overlap checks. They are excluded from all model feature sets.")
    doc.add_heading("C. Class Distribution", level=2)
    add_para(doc, "The cleaned dataset contains 62,500 fraud transactions and 437,485 legitimate transactions, corresponding to a 12.5004% fraud ratio. Fig. 3 summarizes this distribution.")
    add_figure(doc, FIG_DIR / "figure3_class_distribution.png", "Fig. 3. Class distribution of the cleaned synthetic transaction dataset.", width=5.5)
    doc.add_heading("D. Progressive Dataset Versions", level=2)
    add_para(doc, "Five dataset versions are constructed from the same transaction population. Table II summarizes their feature additions, and Fig. 2 visualizes the progressive design.")
    add_table_from_csv(doc, TABLE_DIR / "table2_progressive_dataset_versions.csv", "TABLE II. Progressive Dataset Versions", [900, 1100, 900, 3900, 2560])
    add_figure(doc, FIG_DIR / "figure2_dataset_version_progression.png", "Fig. 2. Progressive dataset version design from v1 through v5.", width=6.0)
    doc.add_heading("E. Feature Exclusion and Leakage Controls", level=2)
    add_para(doc, "The active modeling pipeline excludes FraudFlag, TransactionID, Time, DayOfWeek, Month, IsWeekend, and IsWeekendDerived from model features. Distribution-dependent engineered features are recomputed inside the modeling pipeline from training-fold data. MerchantRisk and CardRisk are treated as synthetic proxy risk indicators and are discussed as leakage-risk features rather than proven deployment-safe variables [25]-[28].")

    doc.add_heading("IV. Methodology", level=1)
    doc.add_heading("A. Experimental Pipeline", level=2)
    add_figure(doc, FIG_DIR / "figure1_experimental_pipeline.png", "Fig. 1. Overall experimental pipeline for dataset construction, model tuning, validation selection, test evaluation, SMOTE comparison, and threshold optimization.", width=5.3)
    add_para(doc, "Fig. 1 summarizes the experimental workflow. The pipeline begins with synthetic dataset construction, cleaning, audit, and five-version dataset generation. Each version is then processed through model-specific preprocessing, cross-validated hyperparameter tuning, validation-based selection, untouched test evaluation, and threshold optimization.")
    add_table_from_csv(doc, TABLE_DIR / "table3_experimental_configuration.csv", "TABLE III. Experimental Configuration", [2500, 6860])
    doc.add_heading("B. Train/Validation/Test Split", level=2)
    add_para(doc, "The experiment uses a stratified 60/20/20 split. The training split contains 299,991 records, the validation split contains 99,997 records, and the test split contains 99,997 records. Hyperparameter search is restricted to the training split, model-family selection is performed on the validation split, and final reporting uses the untouched test split.")
    doc.add_heading("C. Preprocessing Strategy", level=2)
    add_para(doc, "Logistic Regression uses imputation, one-hot encoding for categorical variables, and numeric scaling. Decision Tree, Random Forest, and XGBoost use imputation and frequency encoding for categorical variables. CatBoost follows the repository implementation: leakage-safe feature engineering is applied before CatBoost fitting and prediction, with categorical features handled through the final CatBoost workflow.")
    doc.add_heading("D. Model Selection and Hyperparameter Tuning", level=2)
    add_para(doc, "All model families are tuned with 5-fold cross-validation on the training split using PR-AUC, implemented as average precision. Candidate models are selected on validation fraud-class F1 with tie-breaking based on PR-AUC, recall, precision, ROC-AUC, and lower training time. Randomized hyperparameter search follows standard validation practice [23]-[26].")
    doc.add_heading("E. SMOTE Configuration", level=2)
    add_para(doc, "SMOTE is evaluated as a separate experimental condition. It is applied only inside the imblearn training pipeline, after preprocessing, and inside cross-validation folds. No SMOTE is applied before the train/validation/test split.")
    doc.add_heading("F. Threshold Optimization", level=2)
    add_para(doc, "Threshold optimization is performed after training by selecting a fraud-probability threshold on validation predictions and evaluating that selected threshold on test predictions. This provides an operating-point analysis distinct from default-threshold model comparison.")
    doc.add_heading("G. Evaluation Metrics", level=2)
    add_para(doc, "The study reports fraud-class precision, fraud-class recall, fraud-class F1, ROC-AUC, and PR-AUC. Accuracy is not used as the primary selection criterion because the minority class is the operational focus.")

    doc.add_heading("V. Experimental Results", level=1)
    doc.add_heading("A. Default-Threshold No-SMOTE Results", level=2)
    add_para(doc, "Table IV reports the best default-threshold no-SMOTE model for each dataset version. Performance increases sharply from v1-v2 to v3-v5, indicating that the risk-enriched and engineered versions provide more useful tabular predictors.")
    add_table_from_csv(doc, TABLE_DIR / "table4_best_default_no_smote.csv", "TABLE IV. Best Default-Threshold Results Without SMOTE", [850, 1600, 1300, 1200, 1700, 1200])
    doc.add_heading("B. Default-Threshold SMOTE Results", level=2)
    add_para(doc, "Table V reports the best default-threshold SMOTE model for each dataset version. SMOTE improves fraud-class F1 at the default threshold for all five versions. The best default-threshold result is v3 SMOTE Random Forest with fraud-class F1 of 0.861846 and PR-AUC of 0.855235.")
    add_table_from_csv(doc, TABLE_DIR / "table5_best_default_smote.csv", "TABLE V. Best Default-Threshold Results With SMOTE", [850, 1600, 1300, 1200, 1700, 1200])
    doc.add_heading("C. Feature Engineering Impact", level=2)
    add_para(doc, "Fig. 4 shows the fraud-class F1 comparison across dataset versions. The strongest performance jump occurs when moving from v2 to v3. Versions v3, v4, and v5 are close, so the evidence supports practical similarity among the strongest feature sets rather than a claim that any one of these versions is materially dominant.")
    add_figure(doc, FIG_DIR / "figure4_fraud_f1_comparison.png", "Fig. 4. Fraud-class F1 comparison between no-SMOTE and SMOTE selected models across dataset versions.", width=6.1)
    doc.add_heading("D. SMOTE Impact", level=2)
    add_para(doc, "SMOTE improves default-threshold fraud-class F1 but does not universally improve recall. Fig. 5 and Fig. 6 show that the SMOTE setting often increases precision while reducing recall for the selected default-threshold models. Fig. 7 summarizes the PR-AUC comparison.")
    add_figure(doc, FIG_DIR / "figure5_precision_comparison.png", "Fig. 5. Fraud precision comparison between no-SMOTE and SMOTE selected models.", width=6.1)
    add_figure(doc, FIG_DIR / "figure6_recall_comparison.png", "Fig. 6. Fraud recall comparison between no-SMOTE and SMOTE selected models.", width=6.1)
    add_figure(doc, FIG_DIR / "figure7_pr_auc_comparison.png", "Fig. 7. PR-AUC comparison between no-SMOTE and SMOTE selected models.", width=6.1)
    doc.add_heading("E. Threshold-Tuned Results", level=2)
    add_para(doc, "Table VI reports validation-based threshold-tuned operating points. Threshold tuning largely removes the practical advantage of SMOTE. The v5 no-SMOTE Random Forest operating point at threshold 0.80 records precision 0.999156, recall 0.758000, and fraud-class F1 0.862030.")
    add_table_from_csv(doc, TABLE_DIR / "table6_threshold_tuned_comparison.csv", "TABLE VI. Threshold-Tuned Comparison", [700, 1550, 1250, 1150, 1450, 1200, 1060])
    add_figure(doc, FIG_DIR / "figure8_threshold_optimization_curve.png", "Fig. 8. Threshold optimization curve for the v5 no-SMOTE Random Forest model.", width=6.1)
    add_figure(doc, FIG_DIR / "figure9_best_tuned_confusion_matrix.png", "Fig. 9. Test-set confusion matrix for the v5 no-SMOTE Random Forest model at threshold 0.80.", width=5.7)
    doc.add_heading("F. Best Model Analysis", level=2)
    add_para(doc, "The strongest default-threshold and threshold-tuned results are achieved by tree-based models. Fig. 10 reports impurity-based feature importance for the selected v5 no-SMOTE Random Forest model. These importance values are model-specific and should not be interpreted as causal effects.")
    add_figure(doc, FIG_DIR / "figure10_feature_importance.png", "Fig. 10. Impurity-based feature importance for the selected v5 no-SMOTE Random Forest model.", width=6.0)
    add_table_from_csv(doc, TABLE_DIR / "table7_summary_of_main_findings.csv", "TABLE VII. Summary of Main Findings", [2200, 3800, 3360])

    doc.add_heading("VI. Discussion", level=1)
    doc.add_heading("A. Interpretation of Main Findings", level=2)
    add_para(doc, "The results support the claim that progressive tabular feature engineering improves fraud-class performance on the synthetic transaction dataset. The improvement is most evident when moving from v1-v2 to v3-v5. The strongest versions are practically close, and the manuscript therefore interprets differences among v3-v5 conservatively.")
    doc.add_heading("B. Why Tree-Based Models Performed Better", level=2)
    add_para(doc, "Tree-based methods are well suited to heterogeneous tabular features because they can model nonlinear thresholds and interactions among amount, category, location, transaction timing, and proxy risk variables [18]-[22]. Logistic Regression is retained as a baseline but is not selected as a best model in the final results.")
    doc.add_heading("C. SMOTE Versus Threshold Optimization", level=2)
    add_para(doc, "The default-threshold results suggest that SMOTE improves fraud-class F1. The threshold-tuned results show a more cautious interpretation: once operating thresholds are selected on validation predictions, no-SMOTE and SMOTE settings become practically similar. This indicates that resampling conclusions should be reported with the associated threshold policy.")
    doc.add_heading("D. Practical Implications", level=2)
    add_para(doc, "For fraud analysis workflows, the results emphasize the importance of feature design, validation-based threshold selection, and minority-class metrics. The highest-F1 tuned operating point is highly precise but has moderate recall, meaning that a different threshold could be preferred if the objective were to capture more fraudulent transactions at the cost of additional false positives.")
    doc.add_heading("E. Reproducibility Considerations", level=2)
    add_para(doc, "The repository provides deterministic dataset files, final training outputs, validation and test prediction files, threshold-tuning outputs, and comparison summaries. The final manuscript uses only the active research-clean pipeline and result folders, not stale README instructions or legacy result folders.")

    doc.add_heading("VII. Threats to Validity", level=1)
    doc.add_heading("A. Synthetic Dataset Limitation", level=2)
    add_para(doc, "The dataset is researcher-generated and synthetic. The repository does not document an institutional financial-data source. Therefore, the results should not be interpreted as evidence of performance on operational card transaction streams.")
    doc.add_heading("B. External Validity Limitation", level=2)
    add_para(doc, "No external validation dataset is present. All results are computed from splits of the same cleaned synthetic master dataset. Cross-institution, temporal, or public-benchmark validation remains future work.")
    doc.add_heading("C. Leakage-Risk Features", level=2)
    add_para(doc, "MerchantRisk and CardRisk are proxy risk indicators and may not be leakage-free. Their availability and construction at prediction time must be validated before these features can be considered deployable.")
    doc.add_heading("D. Threshold Dependence", level=2)
    add_para(doc, "The effect of SMOTE depends on the classification threshold. Default-threshold F1 improves under SMOTE, but validation-tuned operating points are practically tied across no-SMOTE and SMOTE settings.")
    doc.add_heading("E. Absence of Confidence Intervals", level=2)
    add_para(doc, "The repository does not contain repeated-seed experiments, confidence intervals, bootstrap uncertainty estimates, or formal significance tests. Small differences among v3-v5 should therefore be interpreted conservatively.")
    doc.add_heading("F. Documentation Drift and Legacy Scripts", level=2)
    add_para(doc, "The repository contains legacy scripts and stale documentation that do not represent the final research-clean workflow. This manuscript cites only the active final scripts, datasets, and result folders.")

    doc.add_heading("VIII. Conclusion and Future Work", level=1)
    add_para(doc, "This paper evaluated progressive feature engineering and imbalanced learning for credit card fraud detection using a synthetic transaction dataset. Five dataset versions and five model families were evaluated under no-SMOTE and SMOTE settings with stratified splitting, cross-validated PR-AUC tuning, validation-based model selection, and untouched test reporting.")
    add_para(doc, "The results show that v3-v5 substantially outperform v1-v2, that tree-based models dominate the strongest results, and that SMOTE improves default-threshold fraud-class F1. However, threshold tuning largely removes the practical advantage of SMOTE. The highest default-threshold fraud-class F1 is 0.861846 for v3 SMOTE Random Forest, while the highest threshold-tuned fraud-class F1 is 0.862030 for v5 no-SMOTE Random Forest.")
    add_para(doc, "Future work should validate the pipeline on external transaction data, document all feature provenance, test the robustness of MerchantRisk and CardRisk, estimate uncertainty across repeated runs, and evaluate operating thresholds under explicit false-positive and false-negative cost assumptions.")

    doc.add_heading("References", level=1)
    for i, (_, text) in enumerate(REFERENCES, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"[{i}] {text}")
        run.font.size = Pt(8.5)

    core = doc.core_properties
    core.title = "Progressive Feature Engineering and Imbalanced Machine Learning for Credit Card Fraud Detection"
    core.subject = "Credit card fraud detection; synthetic transaction dataset; tabular machine learning"
    core.author = "Author Name(s)"
    core.keywords = "credit card fraud detection, imbalanced classification, feature engineering, SMOTE, Random Forest, XGBoost, CatBoost, PR-AUC"

    doc.save(DOCX_PATH)

    readme = """# Publication Package

This folder contains the IEEE-style manuscript package generated from the active research-clean repository evidence.

## Main files

- `manuscript.docx`
- `manuscript.pdf`
- `publication_readiness_report.md`
- `README_publication_package.md`

## Subfolders

- `figures/`: ten publication figures embedded in the manuscript.
- `tables/`: seven CSV tables embedded in the manuscript.
- `references/`: IEEE reference text and BibTeX file.
- `supplementary_materials/`: copied result summaries and feature-importance data.
- `verification/`: render outputs, manifests, repeated embedding checks, and QC report.

## Evidence Scope

The package uses only the active final research-clean datasets, reports, scripts, and result folders listed in the QC report. Legacy outputs are not used as final evidence.
"""
    (PKG / "README_publication_package.md").write_text(readme, encoding="utf-8")
    print(f"Wrote {DOCX_PATH}")
    print(f"Wrote references to {REF_DIR}")


if __name__ == "__main__":
    build_docx()


